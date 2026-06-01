from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\mikic\Documents\New project 2\noesis-commercial-kit")
OUT = ROOT / "exports"

INK = "123C35"
TEAL = "1E6F62"
MUTED = "6D746F"
PAPER = "FBFAF5"
MIST = "DDE9E5"
WINE = "6D3F4B"
GOLD = "B69A62"


def hex_color(value):
    return RGBColor.from_string(value)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="202725"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(9.5)
    run.font.color.rgb = hex_color(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_margins(section):
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)


def base_doc(title, subtitle, meta=None):
    doc = Document()
    set_margins(doc.sections[0])
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = hex_color("202725")
    styles["Heading 1"].font.name = "Georgia"
    styles["Heading 1"].font.size = Pt(20)
    styles["Heading 1"].font.color.rgb = hex_color(INK)
    styles["Heading 2"].font.name = "Georgia"
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].font.color.rgb = hex_color(INK)

    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.text = "Noesis"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.name = "Georgia"
    hp.runs[0].font.size = Pt(11)
    hp.runs[0].font.color.rgb = hex_color(INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Noesis")
    r.font.name = "Georgia"
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = hex_color(INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Arquitectura operativa per empreses")
    r.font.name = "Aptos"
    r.font.size = Pt(9.5)
    r.font.color.rgb = hex_color(MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.font.name = "Georgia"
    r.font.size = Pt(27)
    r.bold = True
    r.font.color.rgb = hex_color(INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(subtitle)
    r.font.name = "Aptos"
    r.font.size = Pt(11)
    r.font.color.rgb = hex_color(MUTED)

    if meta:
      table = doc.add_table(rows=1, cols=len(meta))
      table.alignment = WD_TABLE_ALIGNMENT.CENTER
      table.autofit = True
      for i, (label, value) in enumerate(meta):
          cell = table.cell(0, i)
          shade_cell(cell, PAPER)
          set_cell_text(cell, f"{label}\n{value}", bold=False, color="202725")
      doc.add_paragraph()

    return doc


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.12
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("• ")
    r.font.color.rgb = hex_color(TEAL)
    r.bold = True
    p.add_run(text)


def note_box(doc, title, text):
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    top = table.cell(0, 0)
    shade_cell(top, INK)
    set_cell_text(top, title, bold=True, color="FFFFFF")
    body = table.cell(1, 0)
    shade_cell(body, PAPER)
    set_cell_text(body, text, color="202725")
    doc.add_paragraph()


def save(doc, name):
    OUT.mkdir(parents=True, exist_ok=True)
    path = OUT / name
    doc.save(path)
    return path


def build_pressupost():
    doc = base_doc(
        "Plantilla de pressupost",
        "Proposta d'implementació per sistemes operatius digitals.",
        [("Client", "[EMPRESA]"), ("Projecte", "[PROCÉS]"), ("Data", "[DATA]")],
    )
    note_box(doc, "Enfocament", "La proposta no ven tecnologia per tecnologia. Presenta una millora de procés amb impacte mesurable.")
    h1(doc, "1. Context")
    para(doc, "[EMPRESA] treballa actualment amb [PROCÉS ACTUAL]. Durant la diagnosi inicial hem detectat friccions en [PUNTS CONCRETS].")
    para(doc, "El repte no és només reduir cost, sinó construir un sistema més clar, mesurable i fàcil de mantenir.")
    h1(doc, "2. Objectiu")
    for item in ["reduir seguiment manual", "ordenar informació dispersa", "recuperar temps operatiu", "millorar control i traçabilitat", "mesurar l'impacte"]:
        bullet(doc, item)
    h1(doc, "3. Fases")
    phases = [
        ("Diagnosi operativa", "Entendre el procés real, identificar friccions i definir què s'ha de mesurar."),
        ("Disseny del sistema", "Definir flux, regles, camps, validacions i responsabilitats."),
        ("Implementació", "Construir el sistema i connectar-lo amb el dia a dia de l'empresa."),
        ("Proves i ajustos", "Provar casos reals, corregir friccions i preparar l'equip."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(["Fase", "Objectiu", "Durada"]):
        shade_cell(table.cell(0, i), INK)
        set_cell_text(table.cell(0, i), head, bold=True, color="FFFFFF")
    for phase, objective in phases:
        row = table.add_row().cells
        set_cell_text(row[0], phase, bold=True, color=INK)
        set_cell_text(row[1], objective)
        set_cell_text(row[2], "[X dies]")
    h1(doc, "4. Inversió")
    for item in ["Cost inicial d'implementació: [€]", "Manteniment mensual: [€ / mes]", "Durada mínima recomanada: [X mesos]"]:
        bullet(doc, item)
    h1(doc, "5. Indicadors")
    para(doc, "El projecte es revisarà amb indicadors com hores recuperades, errors evitables, temps de resposta, tasques manuals reduïdes i control del procés.")
    h1(doc, "6. Condicions")
    para(doc, "Validesa del pressupost: [15/30 dies]. Forma de pagament: [a definir]. Qualsevol ampliació d'abast es valorarà per separat.")
    return save(doc, "noesis_plantilla_pressupost.docx")


def build_diagnosi():
    doc = base_doc(
        "Diagnosi operativa",
        "Document per resumir el procés revisat, les friccions i la recomanació inicial.",
        [("Client", "[EMPRESA]"), ("Responsable", "[NOM]"), ("Data", "[DATA]")],
    )
    h1(doc, "1. Procés revisat")
    para(doc, "[Explicar quin procés s'ha revisat: cites, factures, documents, correus, tasques internes, etc.]")
    h1(doc, "2. Situació actual")
    for item in ["Com entra la informació", "Qui la gestiona", "Quines eines s'utilitzen", "On es repeteix feina", "On apareixen errors", "On es perd temps"]:
        bullet(doc, item + ": [detall]")
    h1(doc, "3. Friccions detectades")
    table = doc.add_table(rows=1, cols=3)
    for i, head in enumerate(["Punt", "Observació", "Impacte"]):
        shade_cell(table.cell(0, i), INK)
        set_cell_text(table.cell(0, i), head, bold=True, color="FFFFFF")
    for n in range(1, 4):
        row = table.add_row().cells
        set_cell_text(row[0], str(n), bold=True, color=TEAL)
        set_cell_text(row[1], "[FRICCIÓ]")
        set_cell_text(row[2], "[TEMPS / ERROR / CONTROL]")
    h1(doc, "4. Oportunitat")
    para(doc, "Veiem oportunitat en [PROCÉS] perquè actualment depèn de [seguiment manual / correus / documents / validacions repetides].")
    h1(doc, "5. Recomanació")
    note_box(doc, "Primer pas recomanat", "Començar per [PRIMER SISTEMA], perquè és el punt amb millor relació entre simplicitat d'implementació i impacte.")
    return save(doc, "noesis_plantilla_diagnosi_operativa.docx")


def build_idea():
    doc = base_doc(
        "Idea de servei",
        "Fitxa breu per enviar a un client quan detectem una oportunitat concreta.",
        [("Client", "[EMPRESA]"), ("Servei", "[SISTEMA]"), ("Data", "[DATA]")],
    )
    h1(doc, "Punt observat")
    para(doc, "[EMPRESA] sembla tenir un procés recurrent relacionat amb [PROCÉS].")
    h1(doc, "Possible fricció")
    for item in ["temps dedicat a seguiment manual", "informació repartida entre eines", "errors o duplicats", "poca visibilitat sobre l'estat real", "dependència de persones concretes"]:
        bullet(doc, item)
    h1(doc, "Proposta Noesis")
    para(doc, "Crear un sistema digital que ordeni [PROCÉS] i permeti centralitzar informació, reduir tasques repetitives, definir regles clares i mesurar l'impacte.")
    h1(doc, "Valor per al client")
    para(doc, "El valor no és només fer-ho més ràpid. És que l'empresa pugui treballar amb més control, menys fricció i menys dependència de seguiment manual.")
    h1(doc, "Primer pas")
    note_box(doc, "Diagnosi breu", "Fer una diagnosi operativa per validar si realment hi ha impacte suficient abans de proposar cap implementació.")
    return save(doc, "noesis_plantilla_idea_servei.docx")


def build_normes():
    doc = base_doc(
        "Normes de correu",
        "Guia interna per mantenir una veu comercial coherent, humana i professional.",
        [("Ús", "Intern"), ("Marca", "Noesis"), ("Versió", "1.0")],
    )
    h1(doc, "Principi")
    para(doc, "Cada correu ha de transmetre que Noesis ha mirat el negoci amb criteri i que escriu perquè hi pot haver una millora concreta.")
    h1(doc, "To")
    for item in ["clar", "proper", "professional", "breu", "amb criteri empresarial", "sense llenguatge tècnic innecessari"]:
        bullet(doc, item)
    h1(doc, "Paraules recomanades")
    para(doc, "processos, operativa, fricció, sistema digital, criteri, temps recuperable, ordre, impacte mesurable, seguiment manual, informació dispersa, diagnosi operativa.")
    h1(doc, "Paraules a evitar")
    para(doc, "IA com a reclam principal, revolucionari, màgic, automatització total, garantit, robot, algoritme, eina definitiva.")
    h1(doc, "Estructura del primer correu")
    for item in ["presentació curta", "context concret del negoci", "hipòtesi de fricció", "què pot aportar Noesis", "proposta suau de conversa", "tancament respectuós"]:
        bullet(doc, item)
    return save(doc, "noesis_normes_correu.docx")


if __name__ == "__main__":
    paths = [build_pressupost(), build_diagnosi(), build_idea(), build_normes()]
    for path in paths:
        print(path)
